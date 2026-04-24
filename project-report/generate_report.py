"""
Generate print-ready A4 DOCX and PDF from project-report.md.
Body text: 13pt (PDF) / 14pt (DOCX). Headings scaled for hierarchy.
"""
from pathlib import Path
import re
import textwrap

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, Preformatted, Table, TableStyle
from reportlab.lib import colors

from pypdf import PdfReader, PdfWriter

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "project-report.md"
DELIVERABLES = ROOT / "deliverables"
DELIVERABLES.mkdir(parents=True, exist_ok=True)

DOCX_OUT = DELIVERABLES / "project-report.docx"
PDF_OUT = DELIVERABLES / "project-report.pdf"
PDF_COMBINED = DELIVERABLES / "project-report-complete-with-synopsis.pdf"
SYNOPSIS = DELIVERABLES / "Synopsis Jai.pdf"
DOCX_FALLBACK = DELIVERABLES / "project-report-updated.docx"

# PDF layout — readable for A4 print
PDF_MARGIN = 0.95 * inch
PDF_BODY_SIZE = 13
PDF_BODY_LEADING = 20
PDF_H1_SIZE = 20
PDF_H2_SIZE = 16
PDF_H3_SIZE = 14


def parse_markdown_lines(md_text: str):
    return md_text.splitlines()


def format_inline_for_reportlab(text: str) -> str:
    """Escape XML and convert **bold** to <b></b> for ReportLab Paragraph."""
    s = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
    return s


def strip_markdown_for_docx(text: str) -> str:
    """Remove ** for plain DOCX; keep readable text."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def parse_qr_grid_directive(line: str):
    m = re.match(r"^\{\{QR_GRID:\s*(.*)\s*\}\}$", line.strip())
    if not m:
        return None
    raw = m.group(1)
    entries = []
    for part in raw.split(";;"):
        part = part.strip()
        if not part:
            continue
        bits = [x.strip() for x in part.split("|")]
        if len(bits) != 3:
            continue
        entries.append({"title": bits[0], "image": bits[1], "url": bits[2]})
    return entries if entries else None


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def build_flowable_image(rel: str, alt: str, story, body_style):
    img_path = (ROOT / rel).resolve()
    if not img_path.exists():
        story.append(Paragraph(f"[Figure missing: {rel}]", body_style))
        story.append(Spacer(1, 10))
        return
    max_w = 6.35 * inch
    max_h = 5.05 * inch
    if PILImage is not None:
        try:
            pil = PILImage.open(img_path)
            pw, ph = pil.size
            if pw <= 0 or ph <= 0:
                raise ValueError("bad image size")
            scale = min(max_w / pw, max_h / ph)
            w_in = pw * scale
            h_in = ph * scale
            story.append(Image(str(img_path), width=w_in, height=h_in))
        except Exception:
            story.append(Image(str(img_path), width=max_w, height=3.8 * inch))
    else:
        story.append(Image(str(img_path), width=max_w, height=3.8 * inch))
    cap = format_inline_for_reportlab(f"Figure: {alt}")
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<i>{cap}</i>", body_style))
    story.append(Spacer(1, 12))


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 10)
    canvas.setFillColor(HexColor("#333333"))
    w, h = A4
    if doc.page == 1:
        page_label = ""
    elif doc.page == 2:
        page_label = "Page 0.5"
    else:
        # Report content after Demo Links should align after Synopsis pages 01-37
        # so first chapter page is displayed as 38.
        page_label = f"Page {doc.page + 35}"
    if page_label:
        canvas.drawCentredString(w / 2.0, 0.62 * inch, page_label)
    canvas.setFont("Times-Roman", 8)
    canvas.drawCentredString(w / 2.0, 0.45 * inch, "E-Commerce Website — Academic Project Report")
    canvas.restoreState()


def build_docx(lines):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    image_re = re.compile(r"^!\[(.*?)\]\((.*?)\)$")
    bullet_re = re.compile(r"^\s*-\s+(.*)$")

    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        cols = len(table_rows[0])
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(table_rows[0]):
            hdr[i].text = t
        for row_vals in table_rows[1:]:
            r = table.add_row().cells
            for i, t in enumerate(row_vals):
                r[i].text = t
        doc.add_paragraph("")
        table_rows = []

    skip_thank_you_line = False
    last_was_page_break = False
    for raw in lines:
        line = raw.rstrip()
        qr_entries = parse_qr_grid_directive(line)
        if qr_entries:
            table = doc.add_table(rows=len(qr_entries), cols=1)
            table.style = "Table Grid"
            for idx, entry in enumerate(qr_entries):
                cell = table.cell(idx, 0)
                p1 = cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run(entry["title"])
                r1.bold = True
                r1.font.size = Pt(11)

                img_path = (ROOT / entry["image"]).resolve()
                if img_path.exists():
                    p_img = cell.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(str(img_path), width=Inches(1.4))

                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(entry["url"])
                r2.font.size = Pt(9)
            doc.add_paragraph("")
            continue

        if line.strip().startswith("<!--") and line.strip().endswith("-->"):
            continue
        if line.strip() == r"\newpage":
            if last_was_page_break:
                continue
            flush_table()
            doc.add_page_break()
            last_was_page_break = True
            continue
        last_was_page_break = False

        m_head = heading_re.match(line)
        if m_head:
            flush_table()
            lvl = min(len(m_head.group(1)), 3)
            text = strip_markdown_for_docx(m_head.group(2).strip())
            if lvl == 1 and text.strip().lower() == "the end":
                # Custom final page layout
                p1 = doc.add_paragraph()
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run("The End")
                r1.bold = True
                r1.font.name = "Times New Roman"
                r1.font.size = Pt(42)
                p1.paragraph_format.space_before = Pt(40)
                p1.paragraph_format.space_after = Pt(160)

                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run("Thank you.")
                r2.bold = True
                r2.font.name = "Times New Roman"
                r2.font.size = Pt(34)
                skip_thank_you_line = True
                continue
            h = doc.add_heading(text, level=lvl)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(18 if lvl == 1 else (15 if lvl == 2 else 13))
            continue

        m_img = image_re.match(line.strip())
        if m_img:
            flush_table()
            alt = m_img.group(1).strip()
            rel = m_img.group(2).strip()
            img_path = (ROOT / rel).resolve()
            if img_path.exists():
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.4))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(f"Figure: {alt}")
                    r.italic = True
                    r.font.size = Pt(12)
                except Exception:
                    doc.add_paragraph(f"[Image could not be inserted: {rel}]")
            else:
                doc.add_paragraph(f"[Missing image: {rel}]")
            continue

        if is_table_line(line):
            parts = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(set(p) <= set("-:") for p in parts):
                continue
            table_rows.append(parts)
            continue

        if table_rows and not is_table_line(line):
            flush_table()

        m_bullet = bullet_re.match(line)
        if m_bullet:
            doc.add_paragraph(strip_markdown_for_docx(m_bullet.group(1)), style="List Bullet")
            continue

        if line.strip() == "":
            doc.add_paragraph("")
        else:
            if skip_thank_you_line and line.strip().lower() in {"thank you.", "thank you"}:
                skip_thank_you_line = False
                continue
            doc.add_paragraph(strip_markdown_for_docx(line))

    flush_table()
    try:
        doc.save(DOCX_OUT)
        return DOCX_OUT
    except PermissionError:
        doc.save(DOCX_FALLBACK)
        return DOCX_FALLBACK


def build_pdf(lines):
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=PDF_BODY_SIZE,
        leading=PDF_BODY_LEADING,
        spaceAfter=8,
        alignment=TA_LEFT,
        firstLineIndent=0,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=PDF_H1_SIZE,
        leading=PDF_H1_SIZE + 8,
        spaceAfter=14,
        spaceBefore=6,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=PDF_H2_SIZE,
        leading=PDF_H2_SIZE + 6,
        spaceAfter=10,
        spaceBefore=12,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Times-Bold",
        fontSize=PDF_H3_SIZE,
        leading=PDF_H3_SIZE + 4,
        spaceAfter=8,
        spaceBefore=8,
    )
    end_top = ParagraphStyle(
        "EndTop",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=46,
        leading=52,
        alignment=TA_LEFT,
        spaceAfter=0,
    )
    end_mid = ParagraphStyle(
        "EndMid",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=38,
        leading=44,
        alignment=TA_LEFT,
        spaceAfter=0,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=22,
        bulletIndent=10,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=7.8,
        leading=8.5,
        leftIndent=10,
        rightIndent=4,
        spaceBefore=2,
        spaceAfter=2,
    )

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=PDF_MARGIN,
        rightMargin=PDF_MARGIN,
        topMargin=PDF_MARGIN,
        bottomMargin=1.1 * inch,
        title="E-Commerce Website Project Report",
    )

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    image_re = re.compile(r"^!\[(.*?)\]\((.*?)\)$")
    bullet_re = re.compile(r"^\s*-\s+(.*)$")
    table_re = re.compile(r"^\|.*\|$")

    story = []
    skip_thank_you_line = False
    last_was_page_break = False
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        if line.strip().startswith("<!--") and line.strip().endswith("-->"):
            i += 1
            continue
        qr_entries = parse_qr_grid_directive(line)
        if qr_entries:
            rows = []
            for entry in qr_entries:
                cell_flow = []
                cell_flow.append(Paragraph(f"<b>{format_inline_for_reportlab(entry['title'])}</b>", body))
                cell_flow.append(Spacer(1, 6))
                img_path = (ROOT / entry["image"]).resolve()
                if img_path.exists():
                    cell_flow.append(Image(str(img_path), width=1.35 * inch, height=1.35 * inch))
                cell_flow.append(Spacer(1, 6))
                url_txt = format_inline_for_reportlab(entry["url"])
                cell_flow.append(Paragraph(url_txt, ParagraphStyle("QrUrl", parent=body, fontSize=9.2, leading=11)))
                rows.append([cell_flow])

            usable_width = A4[0] - (PDF_MARGIN * 2)
            t = Table(rows, colWidths=[usable_width])
            t.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.8, colors.HexColor("#222222")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 14))
            i += 1
            continue
        if line.strip() == r"\newpage":
            if last_was_page_break:
                i += 1
                continue
            story.append(PageBreak())
            last_was_page_break = True
            i += 1
            continue
        last_was_page_break = False

        hm = heading_re.match(line)
        if hm:
            lvl = len(hm.group(1))
            text = format_inline_for_reportlab(hm.group(2).strip())
            if lvl == 1 and hm.group(2).strip().lower() == "the end":
                # Custom final page layout for large centered finish text
                end_top_center = ParagraphStyle("EndTopCenter", parent=end_top, alignment=1)
                end_mid_center = ParagraphStyle("EndMidCenter", parent=end_mid, alignment=1)
                story.append(Spacer(1, 60))
                story.append(Paragraph(text, end_top_center))
                story.append(Spacer(1, 220))
                story.append(Paragraph("Thank you.", end_mid_center))
                skip_thank_you_line = True
                i += 1
                continue
            if lvl == 1:
                story.append(Paragraph(text, h1))
            elif lvl == 2:
                story.append(Paragraph(text, h2))
            else:
                story.append(Paragraph(text, h3))
            i += 1
            continue

        im = image_re.match(line.strip())
        if im:
            alt = im.group(1).strip()
            rel = im.group(2).strip()
            build_flowable_image(rel, alt, story, body)
            i += 1
            continue

        bm = bullet_re.match(line)
        if bm:
            t = format_inline_for_reportlab(bm.group(1))
            story.append(Paragraph(f"• {t}", bullet_style))
            i += 1
            continue

        # Compact code block rendering for indented lines
        if line.startswith("    "):
            code_lines = []
            while i < n and lines[i].rstrip().startswith("    "):
                raw_code = lines[i].rstrip()[4:]
                wrapped = textwrap.wrap(
                    raw_code,
                    width=108,
                    break_long_words=True,
                    break_on_hyphens=False,
                    drop_whitespace=False,
                )
                if not wrapped:
                    code_lines.append("")
                elif len(wrapped) == 1:
                    code_lines.append(wrapped[0])
                else:
                    code_lines.append(wrapped[0])
                    for cont in wrapped[1:]:
                        code_lines.append("↳ " + cont)
                i += 1
            story.append(Preformatted("\n".join(code_lines), code_style))
            continue

        if table_re.match(line.strip()):
            t = format_inline_for_reportlab(line.replace("|", "  |  "))
            story.append(Paragraph(t, body))
            i += 1
            continue

        if line.strip() == "":
            story.append(Spacer(1, 6))
        else:
            if skip_thank_you_line and line.strip().lower() in {"thank you.", "thank you"}:
                skip_thank_you_line = False
                i += 1
                continue
            story.append(Paragraph(format_inline_for_reportlab(line), body))
        i += 1

    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def build_combined_pdf():
    if not SYNOPSIS.exists():
        return
    writer = PdfWriter()
    if not PDF_OUT.exists():
        return

    report_reader = PdfReader(str(PDF_OUT))
    synopsis_reader = PdfReader(str(SYNOPSIS))

    # Front pages before synopsis (order in project-report.md):
    # 0) Table of Contents
    # 1) Demo Links (Page 0.5 in footer)
    pre_pages = min(2, len(report_reader.pages))
    for i in range(pre_pages):
        writer.add_page(report_reader.pages[i])

    for page in synopsis_reader.pages:
        writer.add_page(page)

    # Append main report from the first "Chapter 1" page onward. PDF page index 2 is
    # footer "Page 38"; using range(3, ...) was wrong and dropped that page, so the
    # combined file jumped from synopsis page 37 to main content page 39.
    for i in range(2, len(report_reader.pages)):
        writer.add_page(report_reader.pages[i])

    with open(PDF_COMBINED, "wb") as f:
        writer.write(f)


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = parse_markdown_lines(md_text)
    actual_docx = build_docx(lines)
    build_pdf(lines)
    build_combined_pdf()
    print(f"Generated: {actual_docx}")
    print(f"Generated: {PDF_OUT}")
    if PDF_COMBINED.exists():
        print(f"Generated: {PDF_COMBINED}")


if __name__ == "__main__":
    main()
